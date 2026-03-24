#!/usr/bin/env python3
"""Convert a .docx file to Markdown using mammoth + markdownify.

Extracts header/footer content via python-docx and prepends/appends them
to the converted Markdown output. Supports both text and image-based
headers/footers.
"""

import sys
from pathlib import Path

import mammoth
from docx import Document
from markdownify import markdownify as md

# XML namespaces used in OOXML
_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}


def _extract_images_from_part(part, output_dir: Path, prefix: str) -> list[Path]:
    """Extract embedded images from a header/footer part.

    Searches the part's XML for drawing elements with image relationships,
    saves them to output_dir, and returns a list of saved file paths.
    Deduplicates by image blob content.
    """
    saved: list[Path] = []
    seen_blobs: set[int] = set()  # track by blob hash to deduplicate

    el = part._element
    # Find all <a:blip> elements which reference embedded images
    blips = el.findall(".//{%s}blip" % _NS["a"])
    for blip in blips:
        embed_id = blip.get("{%s}embed" % _NS["r"])
        if not embed_id:
            continue
        try:
            img_part = part.part.rels[embed_id].target_part
        except KeyError:
            continue

        blob_hash = hash(img_part.blob)
        if blob_hash in seen_blobs:
            continue
        seen_blobs.add(blob_hash)

        # Determine file extension from content type (e.g. image/jpeg -> .jpeg)
        ext = img_part.content_type.split("/")[-1]
        if ext == "jpeg":
            ext = "jpg"
        img_filename = f"{prefix}_{len(saved) + 1}.{ext}"
        img_path = output_dir / img_filename
        img_path.write_bytes(img_part.blob)
        saved.append(img_path)

    return saved


def _extract_header_footer(input_path: Path, output_dir: Path) -> tuple[str, str]:
    """Extract header and footer content from all sections of a DOCX file.

    Handles both text-based and image-based headers/footers.
    Returns (header_md, footer_md) as Markdown strings.
    Deduplicates identical content across sections.
    """
    doc = Document(str(input_path))
    headers: list[str] = []
    footers: list[str] = []

    for section in doc.sections:
        # --- Headers (prefer first-page header if it exists, else default) ---
        for hdr in (section.first_page_header, section.header):
            if hdr is None:
                continue
            # Try text first
            text = "\n".join(p.text.strip() for p in hdr.paragraphs if p.text.strip())
            if text and text not in headers:
                headers.append(text)
                break

            # Try images if no text found
            imgs = _extract_images_from_part(hdr, output_dir, "header")
            if imgs:
                md_imgs = "  ".join(f"![header]({img.name})" for img in imgs)
                if md_imgs not in headers:
                    headers.append(md_imgs)
                break

        # --- Footers (prefer first-page footer if it exists, else default) ---
        for ftr in (section.first_page_footer, section.footer):
            if ftr is None:
                continue
            # Try text first
            text = "\n".join(p.text.strip() for p in ftr.paragraphs if p.text.strip())
            if text and text not in footers:
                footers.append(text)
                break

            # Try images if no text found
            imgs = _extract_images_from_part(ftr, output_dir, "footer")
            if imgs:
                md_imgs = "  ".join(f"![footer]({img.name})" for img in imgs)
                if md_imgs not in footers:
                    footers.append(md_imgs)
                break

    return "\n\n".join(headers), "\n\n".join(footers)


def convert_docx_to_md(input_path: Path) -> Path:
    """Convert a single .docx file to .md in the same directory."""
    output_path = input_path.with_suffix(".md")

    # Step 1: Extract header/footer via python-docx (text or images)
    header_text, footer_text = _extract_header_footer(input_path, input_path.parent)

    # Step 2: docx -> HTML (mammoth preserves structure: headings, lists, tables, bold/italic)
    with open(input_path, "rb") as f:
        result = mammoth.convert_to_html(f)

    html = result.value

    # Show warnings if any (e.g. unsupported styles)
    if result.messages:
        for msg in result.messages:
            print(f"  ⚠ {msg}", file=sys.stderr)

    # Step 3: HTML -> Markdown (markdownify handles clean conversion)
    markdown = md(html, heading_style="ATX", strip=["img"])

    # Step 4: Assemble final output with header/footer
    parts: list[str] = []

    if header_text:
        parts.append(
            f"> **Header**\n>\n> {header_text.replace(chr(10), chr(10) + '> ')}"
        )
        parts.append("---")
        print(f"  ℹ Header extracted", file=sys.stderr)

    parts.append(markdown)

    if footer_text:
        parts.append("---")
        parts.append(
            f"> **Footer**\n>\n> {footer_text.replace(chr(10), chr(10) + '> ')}"
        )
        print(f"  ℹ Footer extracted", file=sys.stderr)

    # Write output
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(parts))

    return output_path


def main():
    if len(sys.argv) < 2:
        print("Usage: python docx2md.py input.docx [input2.docx ...]")
        sys.exit(1)

    for arg in sys.argv[1:]:
        input_file = Path(arg).expanduser().resolve()

        if not input_file.exists():
            print(f"✗ File not found: {input_file}", file=sys.stderr)
            sys.exit(1)

        if input_file.suffix.lower() != ".docx":
            print(f"✗ Not a .docx file: {input_file}", file=sys.stderr)
            sys.exit(1)

        output_file = convert_docx_to_md(input_file)
        print(f"✓ Converted {input_file} → {output_file}")


if __name__ == "__main__":
    main()
